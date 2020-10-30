import os
import time

from PyQt5.QtWidgets import QMessageBox, QFileDialog
import docx
from docx.shared import Cm
import MainWindow
from main_window_run import my_app_img_dir, my_app_data
import set_page_corpus_connect
from warning import wcalc


def get_threshold():
    base_path = my_app_img_dir
    app_path = my_app_data
    if os.path.isfile(os.path.join(my_app_data, "threshold.ini")):
        with open(os.path.join(my_app_data, "threshold.ini"), "r+") as f:
            r = float(f.readline().replace('\n', ''))
    else:
        r = 1.5
    return r


def replace_str(old_str: str, obj_str: str, paragraphs):
    for p in paragraphs:
        if old_str in p.text:
            inline = p.runs
            for i in inline:
                if old_str in i.text:
                    text = i.text.replace(old_str, obj_str)
                    i.text = text


def save_docx(ui: MainWindow.Ui_MainWindow):
    """
    将缓存中的图片导出为报告
    :param ui:  界面
    :return: None
    """
    base_dir = my_app_img_dir
    my_app_dir = my_app_data
    paras: dict = set_page_corpus_connect.get_all_path()
    threshold: float = get_threshold()
    localtime = time.asctime(time.localtime(time.time()))
    if os.path.isfile(os.path.join(my_app_dir, "corpus.json")) is False:
        QMessageBox.critical(
            ui.statusbar, "corpus.json Not Found", "发生了一些错误\n你貌似没有完成语料设置。"
        )
        print("你貌似没有完成语料设置")
        return
    if os.path.isfile(os.path.join(my_app_dir, "threshold.ini")) is False:
        QMessageBox.critical(
            ui.statusbar, "corpus.json Not Found", "发生了一些错误\n你貌似没有设置数据分析阈值。"
        )
        print("你貌似没有设置数据分析阈值")
        return
    if os.path.isfile(os.path.join(paras["work_path"], "toal.json")) is False:
        QMessageBox.critical(
            ui.statusbar, "toal.json Not Found", "%s\n你貌似没有设置进行预警分析。" % os.path.join(my_app_dir, "toal.json")
        )
        print("你貌似没有设置进行预警分析")
        return
    if (
            os.path.isfile(os.path.join(base_dir, "corpus.png")) and
            os.path.isfile(os.path.join(base_dir, "train_res.png")) and
            os.path.isfile(os.path.join(base_dir, "ana_A.png")) and
            os.path.isfile(os.path.join(base_dir, "ana_B.png")) and
            os.path.isfile(os.path.join(base_dir, "ana_C.png")) and
            os.path.isfile(os.path.join(base_dir, "ana_D.png")) and
            os.path.isfile(os.path.join(base_dir, "warning.png")) and
            os.path.isfile(os.path.join(base_dir, "word_cloud.png"))
    ) is True:

        filename_choose, _ = QFileDialog.getSaveFileName(
            ui.statusbar,
            "Export Docx Report...",
            '.',  # 起始路径
            "Office Document Files (*.docx)"
        )
        if filename_choose == '':
            print("取消")
            return
        print("从缓存加载..")
        # ######################## 开始了docx的处理 #################
        doc = docx.Document(r"Template/Template.docx")
        for i, p in enumerate(doc.paragraphs):
            if '.图1.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "corpus.png"), width=Cm(14.64))  # 在runs的最后一段文字后添加图片
                p.runs[0].add_break()  # 添加一个折行
            elif '.图2.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "train_res.png"), width=Cm(14.64))
                p.runs[0].add_break()  # 添加一个折行
            elif '.图3.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "ana_A.png"), width=Cm(14.64))
                p.runs[0].add_break()  # 添加一个折行
            elif '.图4.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "ana_C.png"), width=Cm(14.64))
                p.runs[0].add_break()  # 添加一个折行
            elif '.图5.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "ana_B.png"), width=Cm(8))
                p.runs[0].add_break()  # 添加一个折行
            elif '.图6.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "ana_D.png"), width=Cm(8))
                p.runs[0].add_break()  # 添加一个折行
            elif '.图7.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "warning.png"), width=Cm(14.64))
                p.runs[0].add_break()  # 添加一个折行
            elif '.图8.' in p.text:
                p.runs[0].add_break()  # 添加一个折行
                p.runs[0].add_picture(os.path.join(base_dir, "word_cloud.png"), width=Cm(14.64))
                p.runs[0].add_break()  # 添加一个折行
        my_w_out: list = wcalc.get_obj(work_path=paras["work_path"])
        res_obj = wcalc.go_for_str(my_w_out)

        replace_str("我是日期", localtime, doc.paragraphs)

        replace_str("我是数据总量", paras["sum"], doc.paragraphs)
        replace_str("我是阈值", str(threshold), doc.paragraphs)
        replace_str("我是数据所见", "%s%s%s%s" % (res_obj[0], res_obj[1], res_obj[2], res_obj[3]), doc.paragraphs)
        replace_str("我是预警提示", "%s" % (res_obj[4]), doc.paragraphs)

        replace_str("我是数据集名", paras["dataset_name"], doc.paragraphs)

        # for p in doc.paragraphs:
        #     if '我是日期' in p.text:
        #         inline = p.runs
        #         for i in inline:
        #             if '我是日期' in i.text:
        #                 text = i.text.replace('我是日期', localtime)
        #                 i.text = text
        # if '我是时间' in p.text:
        #     p.runs[0].text = p.runs[0].text.replace('我是日期', localtime)
        try:
            doc.save(filename_choose)
            QMessageBox.information(ui.statusbar, "完成", "成功完成了生成", QMessageBox.Ok)

        except:
            QMessageBox.critical(
                ui.statusbar, "Error", "发生了一些错误\n请检查是否有足够的权限访问，或目标文件处于是否处于占用状态。"
            )
    else:
        QMessageBox.critical(ui.statusbar, "Error", "文件不完整。")
        print("文件不完整")
